from Bio import SeqIO
from Bio.Seq import Seq
from Bio import pairwise2
from Bio.pairwise2 import format_alignment
from Bio.Align import PairwiseAligner
import os
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
import re
from fpdf import FPDF
import numpy as np
from concurrent.futures import ThreadPoolExecutor
import multiprocessing
from Bio.SeqUtils import GC, molecular_weight
from Bio.SeqUtils.ProtParam import ProteinAnalysis

class SequenceAnalyzer:
    def __init__(self):
        self.reference_seq = None
        self.reference_protein = None
        self.start_seq = None
        self.end_seq = None
        self.quality_threshold = 20
        self.cpu_count = multiprocessing.cpu_count()
        
    def read_ab1(self, file_path):
        """Read an AB1 file and return the sequence record"""
        try:
            record = SeqIO.read(file_path, "abi")
            return record
        except Exception as e:
            raise Exception(f"Error reading AB1 file: {str(e)}")
            
    def read_multiple_formats(self, file_path):
        """Read multiple sequence file formats"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.ab1':
                return SeqIO.read(file_path, "abi")
            elif file_ext in ['.fasta', '.fa', '.fas']:
                return SeqIO.read(file_path, "fasta")
            elif file_ext in ['.fastq', '.fq']:
                return SeqIO.read(file_path, "fastq")
            elif file_ext in ['.gb', '.gbk']:
                return SeqIO.read(file_path, "genbank")
            elif file_ext == '.embl':
                return SeqIO.read(file_path, "embl")
            elif file_ext == '.txt':
                with open(file_path, 'r') as f:
                    content = f.read().strip()
                    clean_content = re.sub(r'[^ATGCN]', '', content.upper())
                    if len(clean_content) > len(content) * 0.8:
                        from Bio.SeqRecord import SeqRecord
                        return SeqRecord(Seq(clean_content), id="sequence")
                    else:
                        raise Exception("File doesn't appear to contain sequence data")
            else:
                raise Exception(f"Unsupported file format: {file_ext}")
        except Exception as e:
            raise Exception(f"Error reading file {file_path}: {str(e)}")

    def set_reference(self, reference_seq):
        """Set the reference sequence and its protein translation"""
        # Clean the reference sequence
        self.reference_seq = self.clean_sequence(reference_seq)
        try:
            self.reference_protein = str(Seq(self.reference_seq).translate())
        except Exception as e:
            raise Exception(f"Error translating reference sequence: {str(e)}")
            
    def calculate_sequence_statistics(self, sequence):
        """Calculate comprehensive sequence statistics"""
        sequence = sequence.upper()
        length = len(sequence)
        
        base_counts = {
            'A': sequence.count('A'),
            'T': sequence.count('T'),
            'G': sequence.count('G'),
            'C': sequence.count('C'),
            'N': sequence.count('N')
        }
        
        base_percentages = {base: (count / length) * 100 for base, count in base_counts.items()}
        gc_content = GC(sequence)
        mol_weight = molecular_weight(sequence, seq_type='DNA')
        
        return {
            'length': length,
            'base_counts': base_counts,
            'base_percentages': base_percentages,
            'gc_content': gc_content,
            'molecular_weight': mol_weight
        }

    def clean_sequence(self, sequence):
        """Clean DNA sequence by removing non-DNA characters and whitespace"""
        # Remove whitespace and convert to uppercase
        sequence = ''.join(sequence.split()).upper()
        # Remove any non-DNA characters (keeping only A, T, G, C, N)
        sequence = re.sub(r'[^ATGCN]', '', sequence)
        return sequence
        
    def quality_filter_sequence(self, record, threshold=None):
        """Filter sequence based on quality scores"""
        if threshold is None:
            threshold = self.quality_threshold
            
        if not hasattr(record, 'letter_annotations') or 'phred_quality' not in record.letter_annotations:
            return str(record.seq)
            
        quality_scores = record.letter_annotations['phred_quality']
        sequence = str(record.seq)
        
        filtered_seq = ""
        for i, (base, quality) in enumerate(zip(sequence, quality_scores)):
            if quality >= threshold:
                filtered_seq += base
            else:
                filtered_seq += 'N'
                
        return filtered_seq

    def find_sequence_boundaries(self, seq_str):
        """Find the best matching region in the sequence"""
        if not self.reference_seq:
            raise Exception("Reference sequence must be set before analysis")
            
        # Try forward orientation first
        alignments = pairwise2.align.localms(seq_str, self.reference_seq, 2, -1, -10, -0.5)
        if alignments:
            best_alignment = alignments[0]
            _, _, score, start, end = best_alignment
            if score > 100:  # Minimum score threshold
                return start, seq_str, end
                
        # Try reverse complement
        seq_str_rc = str(Seq(seq_str).reverse_complement())
        alignments = pairwise2.align.localms(seq_str_rc, self.reference_seq, 2, -1, -10, -0.5)
        if alignments:
            best_alignment = alignments[0]
            _, _, score, start, end = best_alignment
            if score > 100:  # Minimum score threshold
                return start, seq_str_rc, end
                
        raise Exception("Could not find a good matching region in either orientation")

    def trim_sequence(self, seq, record=None):
        """Trim sequence based on best alignment with reference"""
        if not self.reference_seq:
            raise Exception("Reference sequence not set")
            
        seq_str = self.clean_sequence(str(seq))
        
        # Use quality-based trimming if record is provided
        if record and hasattr(record, 'letter_annotations') and 'phred_quality' in record.letter_annotations:
            seq_str = self.advanced_trim_sequence(record)
        
        try:
            start_idx, seq_str, end_idx = self.find_sequence_boundaries(seq_str)
            return seq_str[start_idx:end_idx]
        except Exception as e:
            raise Exception(f"Error trimming sequence: {str(e)}")
            
    def advanced_trim_sequence(self, record, quality_threshold=20, window_size=10):
        """Advanced sequence trimming based on quality scores"""
        sequence = str(record.seq)
        
        if not hasattr(record, 'letter_annotations') or 'phred_quality' not in record.letter_annotations:
            return sequence
            
        quality_scores = record.letter_annotations['phred_quality']
        
        # Find start position
        start_pos = 0
        for i in range(len(quality_scores) - window_size):
            window_avg = np.mean(quality_scores[i:i + window_size])
            if window_avg >= quality_threshold:
                start_pos = i
                break
                
        # Find end position
        end_pos = len(quality_scores)
        for i in range(len(quality_scores) - window_size, window_size, -1):
            window_avg = np.mean(quality_scores[i - window_size:i])
            if window_avg >= quality_threshold:
                end_pos = i
                break
                
        return sequence[start_pos:end_pos]

    def format_alignment_emboss(self, seq1, seq2, aligned_seq1, aligned_seq2, score, filename, protein=False):
        """Format alignment in EMBOSS Water format"""
        # Calculate identity and similarity
        matches = sum(1 for a, b in zip(aligned_seq1, aligned_seq2) if a == b)
        gaps = aligned_seq1.count('-') + aligned_seq2.count('-')
        length = len(aligned_seq1)
        identity = (matches / length) * 100
        similarity = identity  # For DNA sequences, identity equals similarity

        # Format header exactly like in the screenshot
        header = [
            "Analysis Results:",
            "",
            f"=== {'Protein' if protein else 'Nucleotide'} Analysis for {filename} ===",
            "",
            f"# Length: {length}",
            f"# Identity:     {matches}/{length} ({identity:.1f}%)",
            f"# Similarity:   {matches}/{length} ({similarity:.1f}%)",
            f"# Gaps:         {gaps}/{length} ({(gaps/length*100):.1f}%)",
            f"# Score: {score:.1f}",
            "#",
            "#",
            "#======================================",
            ""
        ]

        # Format alignment blocks exactly like in the screenshot
        blocks = []
        block_size = 50
        ref_pos = 1  # Reference starts from 1
        sample_pos = 122  # Sample starts from corresponding position
        
        for i in range(0, len(aligned_seq1), block_size):
            block_end = min(i + block_size, len(aligned_seq1))
            sample_block = aligned_seq1[i:block_end]  # This is from ab1 file
            ref_block = aligned_seq2[i:block_end]     # This is from reference file
            
            # Calculate end positions
            ref_end = ref_pos + len(ref_block.replace('-', '')) - 1
            sample_end = sample_pos + len(sample_block.replace('-', '')) - 1
            
            # Calculate padding to center the sequence
            sequence_start = 25  # Fixed position where sequence should start
            
            # Format numbers with fixed width fields
            ref_start = f"{ref_pos:>3}"
            ref_end = f"{ref_end:>3}"
            sample_start = f"{sample_pos:>3}"
            sample_end = f"{sample_end:>3}"
            
            # Create match/mismatch line
            match_line = ''
            for s1, s2 in zip(ref_block, sample_block):
                if s1 == s2:
                    match_line += '|'
                else:
                    match_line += '*'
            
            # Format exactly like in the screenshot
            blocks.extend([
                f"reference      {ref_start} {ref_block} {ref_end}",
                f"                   {match_line}",
                f"sample         {sample_start} {sample_block} {sample_end}",
                ""
            ])
            
            # Update positions for next block
            ref_pos = int(ref_end) + 1
            sample_pos = int(sample_end) + 1

        return "\n".join(header + blocks)

    def perform_alignment(self, query_seq, ref_seq, filename, protein=False):
        """Perform sequence alignment between query and reference sequences"""
        if not ref_seq:
            raise Exception("Reference sequence not set")
            
        # Clean sequences
        query_seq = self.clean_sequence(str(query_seq))
        ref_seq = self.clean_sequence(str(ref_seq))
        
        if protein:
            # Convert to protein sequences
            try:
                query_protein = str(Seq(query_seq).translate())
                ref_protein = str(Seq(ref_seq).translate())
                alignments = pairwise2.align.globalms(query_protein, ref_protein, 2, -1, -10, -0.5)
            except Exception as e:
                raise Exception(f"Error in protein translation/alignment: {str(e)}")
        else:
            # Nucleotide alignment
            alignments = pairwise2.align.globalms(query_seq, ref_seq, 2, -1, -10, -0.5)
            
        if not alignments:
            raise Exception("No alignments found")
            
        best_alignment = alignments[0]
        aligned_seq1, aligned_seq2, score, start, end = best_alignment
        
        return self.format_alignment_emboss(query_seq, ref_seq, aligned_seq1, aligned_seq2, score, filename, protein)
        
    def parallel_alignment(self, sequences, reference_seq):
        """Perform alignments in parallel"""
        def align_single(seq_data):
            name, sequence = seq_data
            alignments = pairwise2.align.globalms(sequence, reference_seq, 2, -1, -10, -0.5)
            if alignments:
                return name, alignments[0]
            return name, None
            
        with ThreadPoolExecutor(max_workers=self.cpu_count) as executor:
            results = list(executor.map(align_single, sequences.items()))
            
        return dict(results)

    def translate_dna(self, dna_sequence):
        """Translate DNA sequence to protein sequence"""
        try:
            clean_seq = self.clean_sequence(str(dna_sequence))
            return str(Seq(clean_seq).translate())
        except Exception as e:
            raise Exception(f"Error translating DNA sequence: {str(e)}")
            
    def translate_all_frames(self, sequence):
        """Translate sequence in all 6 reading frames"""
        sequence = sequence.upper()
        translations = {}
        
        # Forward frames
        for frame in range(3):
            frame_seq = sequence[frame:]
            if len(frame_seq) >= 3:
                protein = str(Seq(frame_seq).translate())
                translations[f'Frame +{frame + 1}'] = protein
                
        # Reverse frames
        rev_seq = str(Seq(sequence).reverse_complement())
        for frame in range(3):
            frame_seq = rev_seq[frame:]
            if len(frame_seq) >= 3:
                protein = str(Seq(frame_seq).translate())
                translations[f'Frame -{frame + 1}'] = protein
                
        return translations
        
    def find_restriction_sites(self, sequence, enzymes=None):
        """Find restriction enzyme sites"""
        if enzymes is None:
            enzymes = {
                'EcoRI': 'GAATTC',
                'BamHI': 'GGATCC',
                'HindIII': 'AAGCTT',
                'XbaI': 'TCTAGA',
                'SacI': 'GAGCTC',
                'KpnI': 'GGTACC',
                'SmaI': 'CCCGGG',
                'PstI': 'CTGCAG'
            }
            
        sites = {}
        sequence = sequence.upper()
        
        for enzyme, site in enzymes.items():
            positions = []
            start = 0
            while True:
                pos = sequence.find(site, start)
                if pos == -1:
                    break
                positions.append(pos + 1)
                start = pos + 1
            sites[enzyme] = positions
            
        return sites

    def save_to_word(self, results, filename):
        """Save results to Word document with proper formatting"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_LINE_SPACING

        doc = Document()
        
        # Set font for the entire document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(10)
        
        # Set page width to reasonable size (8.5 inches is standard letter width)
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(8.5)  # Standard letter width
            section.page_height = Inches(11)  # Standard letter height
            # Set margins
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
        
        # Add results exactly as they appear on screen
        for analysis_type, samples in results.items():
            # Add analysis type header
            header = doc.add_paragraph()
            header.add_run(f"\n=== {analysis_type.capitalize()} Analysis Results ===\n\n")
            
            for sample, alignment in samples.items():
                # Add the alignment text directly
                paragraph = doc.add_paragraph()
                # Set line spacing to exactly to prevent auto-adjusting
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(12)
                
                # Clear any existing tab stops
                paragraph.paragraph_format.tab_stops.clear_all()
                
                # Add tab stops every 0.5 inches up to 8.5 inches
                for i in range(1, 17):  # 16 stops (8.5 inches * 2 for half-inch increments)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(i * 0.5), WD_TAB_ALIGNMENT.LEFT)
                
                run = paragraph.add_run(alignment)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                
                # Add page break between samples
                if not (analysis_type == list(results.keys())[-1] and sample == list(samples.keys())[-1]):
                    doc.add_page_break()
        
        # Save document
        doc.save(filename)

    def export_to_pdf(self, results, output_path):
        """Export analysis results to PDF"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Courier", size=10)
        
        # Add title
        pdf.set_font("Courier", 'B', 14)
        pdf.cell(0, 10, "SeqAnalyse Results", ln=True, align='C')
        pdf.ln(10)
        
        # Add results
        pdf.set_font("Courier", size=10)
        for analysis_type, samples in results.items():
            pdf.set_font("Courier", 'B', 12)
            pdf.cell(0, 10, f"{analysis_type.capitalize()} Analysis Results:", ln=True)
            pdf.ln(5)
            
            pdf.set_font("Courier", size=10)
            for sample_name, result in samples.items():
                # Split result into lines
                lines = result.split('\n')
                for line in lines:
                    pdf.cell(0, 5, line, ln=True)
                pdf.ln(5)
            
            pdf.ln(10)
        
        # Save PDF
        pdf.output(output_path)
        
    def export_advanced_results(self, results, output_path, format='pdf'):
        """Export results with advanced formatting"""
        if format.lower() == 'pdf':
            self._export_advanced_pdf(results, output_path)
        elif format.lower() == 'html':
            self._export_html(results, output_path)
        else:
            raise Exception(f"Unsupported export format: {format}")

    def _export_advanced_pdf(self, results, output_path):
        """Export to PDF with advanced formatting"""
        class AdvancedPDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 15)
                self.cell(0, 10, 'SeqAnalyse - Advanced Analysis Results', 0, 1, 'C')
                self.ln(10)
                
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
                
        pdf = AdvancedPDF()
        pdf.add_page()
        
        for analysis_type, samples in results.items():
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, f'{analysis_type.capitalize()} Analysis', 0, 1)
            pdf.ln(5)
            
            for sample_name, result in samples.items():
                pdf.set_font('Arial', 'B', 12)
                pdf.cell(0, 8, f'Sample: {sample_name}', 0, 1)
                
                pdf.set_font('Courier', '', 8)
                lines = result.split('\n')
                for line in lines:
                    if len(line) > 80:
                        for i in range(0, len(line), 80):
                            pdf.cell(0, 4, line[i:i+80], 0, 1)
                    else:
                        pdf.cell(0, 4, line, 0, 1)
                pdf.ln(5)
                
        pdf.output(output_path)

    def _export_html(self, results, output_path):
        """Export results to HTML format"""
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>SeqAnalyse Results</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .header { background-color: #f0f0f0; padding: 10px; text-align: center; }
                .analysis-section { margin: 20px 0; border: 1px solid #ccc; padding: 15px; }
                .sample { margin: 10px 0; }
                .sequence { font-family: 'Courier New', monospace; font-size: 12px; 
                           background-color: #f9f9f9; padding: 10px; white-space: pre-wrap; }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>SeqAnalyse - Advanced Analysis Results</h1>
            </div>
        """
        
        for analysis_type, samples in results.items():
            html_content += f'<div class="analysis-section"><h2>{analysis_type.capitalize()} Analysis</h2>'
            
            for sample_name, result in samples.items():
                html_content += f'<div class="sample"><h3>Sample: {sample_name}</h3>'
                html_content += f'<div class="sequence">{result}</div></div>'
                
            html_content += '</div>'
            
        html_content += '</body></html>'
        
        with open(output_path, 'w') as f:
            f.write(html_content)